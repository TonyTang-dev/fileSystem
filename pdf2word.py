# from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from docx import Document
document = Document()
import warnings
warnings.filterwarnings("ignore")
import os


file_name=os.open('c:\\Users\\Administrator\\Desktop\\22.pdf',os.O_RDWR )

def main():

    fn = open(file_name,'rb')
    parser = PDFParser(fn)
    doc = PDFDocument(parser)
    parser.set_document(doc)
    # doc.set_parser(parser)
    resource = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(resource,laparams=laparams)
    interpreter = PDFPageInterpreter(resource,device)
    for i in PDFPage.create_pages(doc):
        interpreter.process_page(i)
        layout = device.get_result()
        for out in layout:
            if hasattr(out,"get_text"):
                content = out.get_text().replace(u'\xa0', u' ') 
                document.add_paragraph(
                    content, style='ListBullet'   
                )
            document.save('c:\\Users\\Administrator\\Desktop\\b'+'.docx')
    print('处理完成')
 
if __name__ == '__main__':
    main()
    print("a")